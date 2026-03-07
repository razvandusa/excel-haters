import os
import io
import re
import json
import base64

import httpx
import openpyxl
import docx
from dotenv import load_dotenv
from fastapi import FastAPI, UploadFile, File, HTTPException
from fastapi.responses import JSONResponse

# Reads HF_API_KEY from .env file in the same folder as this script
load_dotenv()

#Configuration

HF_API_KEY = os.environ.get("HF_API_KEY", "")
HF_MODEL = "Qwen/Qwen2.5-VL-7B-Instruct"
HF_API_URL = "https://router.huggingface.co/v1/chat/completions"

#App

app = FastAPI()

# Supported MIME types
IMAGE_TYPES = {"image/jpeg", "image/png", "image/webp"}
PDF_TYPE    = "application/pdf"
EXCEL_TYPES = {
    "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
    "application/vnd.ms-excel",
}
WORD_TYPES  = {
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    "application/msword",
}
TEXT_TYPES  = {"text/plain", "text/csv"}

# Prompt for the model — instructs it to extract flight data into a strict JSON format.

EXTRACTION_PROMPT = """
You are a flight data extraction assistant.
Extract ALL flight records from the provided document.
Return ONLY a valid raw JSON object — no explanation, no markdown, no code fences.

Rules:
- Every flight MUST have a flightId.
- arrivalTime and departureTime are mutually exclusive: a flight has one OR the other, never both. Set the missing one to null.
- All times must be ISO 8601 format: YYYY-MM-DDTHH:MM:SS. If no date is in the document use 1970-01-01.
- Optional fields (terminal, desk, security, gate, stand) must be null if not found — never omit them.
- If no flights are found, return: {"flights": []}

Required JSON structure:
{
  "flights": [
    {
      "flightId": "string",
      "terminal": "string or null",
      "desk": "string or null",
      "security": "string or null",
      "gate": "string or null",
      "stand": "string or null",
      "arrivalTime": "ISO8601 string or null",
      "departureTime": "ISO8601 string or null"
    }
  ]
}
"""

# Document processing functions

def image_to_base64_uri(data: bytes, media_type: str) -> str:
    # Convert raw image bytes to a base64 data URI.
    b64 = base64.b64encode(data).decode("utf-8")
    return f"data:{media_type};base64,{b64}"


def pdf_to_image_uris(data: bytes) -> list[str]:
    # Convert each PDF page to a base64 PNG data URI.
    try:
        from pdf2image import convert_from_bytes
        pages = convert_from_bytes(data, dpi=200)
        uris = []
        for page in pages:
            buf = io.BytesIO()
            page.save(buf, format="PNG")
            b64 = base64.b64encode(buf.getvalue()).decode("utf-8")
            uris.append(f"data:image/png;base64,{b64}")
        return uris
    except ImportError:
        return []  # falls back to text extraction below


def extract_text_from_excel(data: bytes) -> str:
    # Read all Excel sheets and return as tab-separated plain text.
    wb = openpyxl.load_workbook(io.BytesIO(data), data_only=True)
    lines = []
    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        lines.append(f"[Sheet: {sheet_name}]")
        for row in ws.iter_rows(values_only=True):
            if any(cell is not None for cell in row):
                lines.append("\t".join(str(c) if c is not None else "" for c in row))
    return "\n".join(lines)


def extract_text_from_word(data: bytes) -> str:
    # Extract all paragraphs and table content from a Word document.
    document = docx.Document(io.BytesIO(data))
    lines = []
    for para in document.paragraphs:
        if para.text.strip():
            lines.append(para.text.strip())
    for table in document.tables:
        for row in table.rows:
            row_text = "\t".join(cell.text.strip() for cell in row.cells)
            if row_text.strip():
                lines.append(row_text)
    return "\n".join(lines)

# HuggingFace API

def call_model_with_images(image_uris: list[str]) -> str:
    # Send images to Qwen2.5-VL and return the raw text response.
    # Build content: one block per image + the prompt at the end
    content = [{"type": "image_url", "image_url": {"url": uri}} for uri in image_uris]
    content.append({"type": "text", "text": EXTRACTION_PROMPT})

    payload = {
        "model": HF_MODEL,
        "messages": [{"role": "user", "content": content}],
        "max_tokens": 4096,
        "temperature": 0.0,  # deterministic — we want consistent JSON every time
    }

    with httpx.Client(timeout=120.0) as client:
        response = client.post(
            HF_API_URL,
            json=payload,
            headers={"Authorization": f"Bearer {HF_API_KEY}"}
        )

    if response.status_code != 200:
        raise ValueError(f"HuggingFace API error {response.status_code}: {response.text[:300]}")

    return response.json()["choices"][0]["message"]["content"]


def call_model_with_text(text: str) -> str:
    # Send plain text to Qwen2.5-VL and return the raw text response.
    payload = {
        "model": HF_MODEL,
        "messages": [
            {
                "role": "user",
                "content": f"{EXTRACTION_PROMPT}\n\nDocument content:\n{text}"
            }
        ],
        "max_tokens": 4096,
        "temperature": 0.0,
    }

    with httpx.Client(timeout=120.0) as client:
        response = client.post(
            HF_API_URL,
            json=payload,
            headers={"Authorization": f"Bearer {HF_API_KEY}"}
        )

    if response.status_code != 200:
        raise ValueError(f"HuggingFace API error {response.status_code}: {response.text[:300]}")

    return response.json()["choices"][0]["message"]["content"]


def parse_response(raw: str) -> dict:
    """Parse the model response — strip markdown fences if present, then JSON parse."""
    clean = re.sub(r"```(?:json)?", "", raw).strip().rstrip("`")
    try:
        return json.loads(clean)
    except json.JSONDecodeError as e:
        raise ValueError(f"Model returned invalid JSON: {e}\nRaw response: {raw[:300]}")

# Endpoint 

@app.post("/flights/import", summary="Extract flight data from any document")
async def extract_flights(file: UploadFile = File(...)):
    """
    Upload any document containing flight data.
    Supported: JPG, PNG, WEBP, PDF, XLSX, DOCX, CSV, TXT
    Returns structured JSON with all detected flights.
    """
    if not HF_API_KEY:
        raise HTTPException(
            status_code=500,
            detail="HF_API_KEY is not set. Add it to your .env file."
        )

    content_type = (file.content_type or "").lower()
    data = await file.read()

    try:
        if content_type in IMAGE_TYPES:
            # Image → send directly to vision model
            uri = image_to_base64_uri(data, content_type)
            raw = call_model_with_images([uri])

        elif content_type == PDF_TYPE:
            # PDF → convert pages to images → send to vision model
            uris = pdf_to_image_uris(data)
            if uris:
                raw = call_model_with_images(uris)
            else:
                # Fallback: extract text with pypdf if pdf2image not installed
                try:
                    import pypdf
                    reader = pypdf.PdfReader(io.BytesIO(data))
                    text = "\n".join(page.extract_text() or "" for page in reader.pages)
                    raw = call_model_with_text(text)
                except ImportError:
                    raise HTTPException(
                        status_code=500,
                        detail="PDF support requires: pip install pdf2image  or  pip install pypdf"
                    )

        elif content_type in EXCEL_TYPES:
            # Excel → extract all sheets as text → send to model
            text = extract_text_from_excel(data)
            raw = call_model_with_text(text)

        elif content_type in WORD_TYPES:
            # Word → extract paragraphs + tables as text → send to model
            text = extract_text_from_word(data)
            raw = call_model_with_text(text)

        elif content_type in TEXT_TYPES:
            # Plain text / CSV → send directly
            text = data.decode("utf-8", errors="replace")
            raw = call_model_with_text(text)

        else:
            raise HTTPException(
                status_code=415,
                detail=f"Unsupported file type: '{content_type}'. "
                       f"Supported: JPG, PNG, PDF, XLSX, DOCX, CSV, TXT."
            )

        result = parse_response(raw)
        return JSONResponse(content=result)

    except HTTPException:
        raise
    except ValueError as e:
        return JSONResponse(content={"error": str(e)}, status_code=422)
    except Exception as e:
        return JSONResponse(content={"error": str(e)}, status_code=500)


@app.get("/health")
def health():
    return {"status": "ok", "model": HF_MODEL}